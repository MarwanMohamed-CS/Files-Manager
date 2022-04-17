import os
import collections
import re
import colorama
import prettytable
import shutil
import datetime
import stat
from time import sleep
from docx import Document
from docx.shared import Inches, Cm, RGBColor
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH
from guessit import guessit
from files_paths import settings


class File():
    '''Carries info of a file and has file management methods'''

    def __init__(self, path, file_type, dir_tag='', size=0, epoch_creation_date=0, surface=False, hard_disk='',):
        self.name = os.path.basename(path)
        self.path = path
        self.dir_tag = dir_tag
        self.file_type = file_type
        self.size = size
        self.surface = surface
        self.hard_disk = hard_disk
        self.epoch_creation_date = epoch_creation_date

    def get_size(self):
        '''returns size(int or float) of file and unit(str) of that size'''
        # size is in bytes
        # 10^3 > KB
        # 10^6 > MB
        if self.size < 10**6:  # converts bytes to KBs
            size = self.size/1024
            unit = 'KB'
            size = int(size)
        elif self.size < 10**9:  # converts bytes to MBs
            # 1048576 is 1024 * 1024
            size = self.size//1_048_576
            unit = 'MB'
        else:  # if size>10^9 converts to GBs
            # 1073741824 is 1024 * 1024 * 1024
            size = self.size/1_073_741_824
            size = round(size, 1)  # rounds to a total of 4 digits
            unit = 'GB'
        if int(size) == 0:
            size = 0
        return size, unit

    def set_index(self, index):
        '''
        sets index for the file so that we can choose it for operations like opening,
        printing tree...etc
        '''
        self.index = index

    def get_creation_date(self):
        '''
        returns a nametuple with the date and the time of the day of the creation of the file
        '''
        datetime_time = datetime.datetime.fromtimestamp(
            float(self.epoch_creation_date))
        datetime_time = str(datetime_time)
        formated_date_in_days, formated_date_in_hm = datetime_time.split(
            ' ')   # hm stands for hours/mins
        #date_formater = collections.namedtuple('creation_date', 'date time')
        return formated_date_in_days, formated_date_in_hm


def remove_ext(file_name):
    '''
    returns file name without it's extension
    '''
    for index, letter in enumerate(file_name[::-1]):
        if letter == '.':
            return file_name[: -index - 1]


def to_folder(file_path, folder_name=''):
    '''
    sends file_path to a folder of it's own name or to a folder with folder_name if it's passed
    '''
    if folder_name:
        folder_path = os.path.join(os.path.dirname(
            remove_ext(file_path)), folder_name)
    else:
        folder_path = remove_ext(file_path)
    new_file_path = folder_path + '\\' + os.path.basename(file_path)
    os.mkdir(folder_path)
    shutil.copy(file_path, new_file_path)
    os.remove(file_path)


def binary_search(paths, search_for):
    searchForWords = search_for.split()
    start = 0
    end = len(paths) - 1
    while start <= end:
        middle_index = int((start + end) / 2)
        midpoint = paths[middle_index].name
        if midpoint > searchForWords[0]:
            end = middle_index - 1
        elif midpoint < searchForWords[0]:
            start = middle_index + 1
        else:
            print(midpoint)
    input()


def organize_files(paths):
    for path in paths:
        if path.surface and path.file_type == 'movie' and os.path.isfile(path.path) and path.dir_tag == 'this pc e':
            input(path.path)
            to_folder(path.path)
            print('waiting')


def parse_inp(inp):
    '''
    parses the input and returns a dictionary containing the corresponding properties to the cmd found
    '''
    net_inp, cmd = get_cmd(inp)
    net_inp, dir_tag = get_dir_tag(inp)
    if cmd:
        cmd['type'] = 'cmd'
    elif dir_tag:
        cmd['type'] = 'dir tag'
        cmd['dir tag'] = dir_tag
    return net_inp, cmd


def get_cmd(inp):
    '''
    returns a dictionary containing the properties corresponding to the inp passed
    '''
    def assign_prop(cmd, cmd_tag):
        '''assigns the properties to the cmd dictionary depedning on the cmd_tag given'''
        if cmd_tag == '-dups':  # to change what function catches as commands
            cmd['func name'] = 'print_duplicates'
        elif cmd_tag == '-all movies':
            cmd['func name'] = 'get_clips'
            cmd['file type'] = 'movie'
            cmd['hard disk'] = []
        elif cmd_tag == '-all tv':
            cmd['func name'] = 'get_clips'
            cmd['file type'] = 'tv series'
        elif cmd_tag == '-all animes':
            cmd['func name'] = 'get_clips'
            cmd['file type'] = 'anime'
        elif cmd_tag == '-pc e':
            cmd['func name'] = 'get_clips'
            cmd['dir tag'] = 'this pc e'
        elif cmd_tag == '-pc f':
            cmd['func name'] = 'get_clips'
            cmd['dir tag'] = 'this pc f'
        elif cmd_tag == '-pc tv':
            cmd['func name'] = 'get_clips'
            cmd['file type'] = 'tv series'
            cmd['dir tag'] = 'this pc'
        elif cmd_tag == '-pc movies':
            cmd['func name'] = 'get_clips'
            cmd['file type'] = 'movie'
            cmd['dir tag'] = 'this pc'
        elif cmd_tag == '-2 tera movies':
            cmd['func name'] = 'get_clips'
            cmd['file type'] = 'movie'
            cmd['dir tag'] = '2 tera movies'
        elif cmd_tag == '-2 tera tv':
            cmd['func name'] = 'get_clips'
            cmd['file type'] = 'tv series'
            cmd['dir tag'] = '2 tera tv'
        elif cmd_tag == '-2 tera animes':
            cmd['func name'] = 'get_clips'
            cmd['file type'] = 'anime'
            cmd['dir tag'] = '2 tera animes'
        elif cmd_tag == '-1 tera movies':
            cmd['func name'] = 'get_clips'
            cmd['file type'] = 'movie'
            cmd['dir tag'] = '1 tera movies'
        elif cmd_tag == '-1 tera tv':
            cmd['func name'] = 'get_clips'
            cmd['file type'] = 'tv series'
            cmd['dir tag'] = '1 tera tv'
        elif cmd_tag == '-1 tera animes':
            cmd['func name'] = 'get_clips'
            cmd['file type'] = 'anime'
            cmd['dir tag'] = '1 tera animes'
        elif cmd_tag == '-u' or cmd_tag == '-update':
            cmd['func name'] = 'update'
        elif cmd_tag == '-new':
            cmd['func name'] = 'get_clips'
            cmd['file type'] = 'new'
    match = re.search('-.*', inp)
    cmd = {}
    if match:
        cmd_tag = match.group()
        net_inp = inp.replace(cmd_tag, '')
    else:
        return inp, cmd
    assign_prop(cmd, cmd_tag)
    return net_inp, cmd


def get_dir_tag(inp):
    '''returns directory tag(if exists) and input line without the tag in it'''
    match = re.search(r'''^         ((this \s pc (\s(e|f))? (?=\s))
                        # this pc followed by e or f opt.,followed by space(not consumed)
                                                        |
                                ((1|2) \s tera (\s (movies|tv|animes) )? (?=\s)))''',
                      # 1|2 tera followed by movies tv or animes(opt.) followed by space(not consumed)
                      inp,
                      re.VERBOSE | re.I)
    if match:
        dir_tag = match.group()
        to_replace = dir_tag+' '  # because dir_tag doesn't have the space in it
        net_inp = inp.replace(to_replace, '')
    else:
        dir_tag = None
        net_inp = inp
    return net_inp, dir_tag


def search(paths, search_for, dir_tag=''):
    '''returns a list of File class objects that carry the results's data/details.'''
    os.system('cls')
    search_for = search_for.lower().split()
    SEARCH_WORDS_COUNT = len(search_for)
    results = []
    index = 0
    for path in paths:
        matches_count = 0
        if path.surface and search_for[0] in path.name.lower():
            for word in search_for:
                if word in path.name.lower():
                    matches_count += 1
            if matches_count == SEARCH_WORDS_COUNT:
                if dir_tag:  # i think if you remove this
                    if dir_tag in path.dir_tag.lower():
                        path.set_index(index)
                        results.append(path)
                        index += 1
                else:  # and this the code will work thus fine because '' is in any string anyway
                    path.set_index(index)
                    results.append(path)
                    index += 1
    return results


def delete(file_path):
    os.chmod(file_path, stat.S_IWRITE)
    os.remove(file_path)


def print_results(results, dups=False):
    '''
        takes in results of a search (File objects or Dup objects) and prints their info in a formatted
        manner
        '''
    colorama.init()
    os.system('cls')
    print(colorama.Fore.YELLOW+'\n\n\t\t\t\t\t\t\tResults:' +
          ' ' * 2 + str(len(results)) + '\n\n\n')
    if not dups:
        if len(results) != 0:
            table = prettytable.PrettyTable([colorama.Fore.YELLOW+'N',
                                             colorama.Fore.YELLOW+'Names',
                                             colorama.Fore.YELLOW+'Dir Tag',
                                             colorama.Fore.YELLOW+'Size',
                                             colorama.Fore.YELLOW+'Creation Date'])
            num = 1

            # if the end results contains the size of of all the results
            if type(results[-1]) == int:
                size_table = prettytable.PrettyTable([colorama.Fore.YELLOW+'Totale Size',
                                                      ])
                totale_results_size = results[-1]
                totale_results_size, unit = to_unit(totale_results_size)
                size_table.add_row([
                    colorama.Fore.LIGHTBLUE_EX+f'{str(totale_results_size)} {unit}'])
                end = len(results) - 2
            else:
                end = len(results)

            for result in results[:end]:
                size, unit = result.get_size()
                if num == settings.RESULTS_LIMIT + 1:
                    break
                date, time = result.get_creation_date()
                table.add_row([colorama.Fore.LIGHTBLUE_EX+f'{num}',
                               colorama.Fore.LIGHTBLUE_EX+f'{result.name:80}',
                               colorama.Fore.LIGHTBLUE_EX +
                               f'{result.dir_tag.title():<10}',
                               colorama.Fore.LIGHTBLUE_EX +
                               f'{size:<4} {unit}',
                               colorama.Fore.LIGHTBLUE_EX + f'{date}'])
                num += 1
            # if the end results contains the size of of all the results
            if type(results[-1]) == int:
                print(table)
                print(size_table)
            else:
                print(table)
        else:
            print(' ' * 51 + '----No results----')
    else:
        for result in results:
            if len(result.dir_tags) > 4:  # prints only 3 and .... the rest
                dir_tags = ', '.join(results.dir_tags[:3])
                dir_tags += ('....')
            else:
                dir_tags = ', '.join(result.dir_tags)
            print(colorama.Fore.BLUE +
                  f'{result.name:70} {colorama.Fore.RED}{dir_tags.title():50}\n')


def get_duplicates(clips):
    '''
    returns a list of Dup namedtuples that contain the info of the duplicate
    surface files that exist in trees
    '''
    print('Guessing names of files..')
    clips_names = [guessit(clip.name)['title']
                   for clip in clips if clip.surface]
    files_count = collections.Counter(clips_names)
    duplicates = []
    dups_names = []
    for file_name, count in files_count.items():
        if count > 1:
            dups_names.append(file_name)
    for dup_name in dups_names:
        results = search(clips, dup_name)  # all results are of same name
        # but they have diff dir_tags/locations
        locations = [result.path for result in results]
        dir_tags = [result.dir_tag for result in results]
        file_type = results[0].file_type
        Dup = collections.namedtuple(
            'dup', 'name file_type locations dir_tags')
        dup_obj = Dup(dup_name, file_type, locations, dir_tags)
        duplicates.append(dup_obj)
    return duplicates


def to_unit(size):
    '''returns size as float\int and an str containing size Unit'''
    # size is in bytes
    # 10^3 > KB
    # 10^6 > MB
    if size < 10**6:  # converts bytes to KBs
        size = size/1024
        size = int(size)
        unit = 'KB'
    elif size < 10**9:  # converts bytes to MBs
        # 1048576 is 1024 * 1024
        size = size/1024**2
        size = int(size)
        unit = 'MB'
    elif size < 10**12:  # converts to GBs
        size = size / 1024**3
        size = round(size, 1)
        unit = 'GB'
    else:  # converts to TBs
        size = size/1024**4
        size = round(size, 3)  # rounds to a totale of 3 digits
        unit = 'TB'
    size = str(size)
    return size, unit


def print_tree(paths):
    '''prints tree with indentation that indicate which folder/file is inside which'''
    os.system('cls')
    cont_path = paths[0]
    print(f'{colorama.Fore.YELLOW}\n\n\t\t\t\t\t\t\tResults\n\n\n')
    cont_size, cont_unit = cont_path.get_size()
    print(colorama.Fore.YELLOW +
          f'\n\n{cont_path.name:70}', end='')  # has no indentation
    print(colorama.Fore.RED +
          f'{cont_size:12} {cont_unit}\n')
    print(colorama.Fore.YELLOW + '-' * 90 + '\n')
    for path in paths[1:]:
        # you can simply know the level by the number of backslashes in the name
        # we remove the backslashes in start to restore the level to 0 for the levels
        # to be in respect to the start_path itself not C:\ or F:\
        level = path.path.replace(cont_path.path, '').count(os.sep) - 1
        indent = ' ' * 8 * level
        size, unit = path.get_size()
        print(colorama.Fore.WHITE +
              f'{indent}{path.name:70}   {colorama.Fore.BLUE}{size:>6} {unit:>}\n')
    input()


def get_clips(paths, cmd):
    '''
    returns a chunk of clips according to dir tag and file type of the Cmd Object(ex:-this pc movies) totale size of results(in bytes) is appended
    to the back of the list. 
    '''
    clips = []
    index = 0
    results_size = 0
    if 'file type' in cmd.keys():
        if 'dir tag' in cmd.keys():
            for path in paths:
                if path.surface and cmd['dir tag'] == path.dir_tag.lower() and cmd['file type'] == path.file_type:
                    path.set_index(index)
                    clips.append(path)
                    results_size += path.size
                    index += 1
        elif cmd['file type'] == 'new':
            creation_ordered_paths = paths[:]
            creation_ordered_paths.sort(
                key=lambda x: x.epoch_creation_date, reverse=True)
            for path in creation_ordered_paths:
                if path.surface:
                    clips.append(path)
                    results_size += path.size
                    index += 1
        else:
            for path in paths:
                if path.surface and cmd['file type'] == path.file_type:
                    path.set_index(index)
                    clips.append(path)
                    results_size += path.size
                    index += 1
    else:
        for path in paths:
            if path.surface and cmd['dir tag'] == path.dir_tag:
                path.set_index(index)
                clips.append(path)
                results_size += path.size
                index += 1
    clips.append(results_size)
    return clips


def get_tree(paths, results, end_cmd):
    '''
    returns a list of the tree of the current result's file path
                    (actual full path is used to search)
    '''
    match = re.search(r'tree(\s\d{1,})?', end_cmd)
    if match:
        try:
            inp_num = match.group().split(' ')[1]
            inp_num = int(inp_num)
        except:
            inp_num = None
    else:
        print('something wrong with the cmd..')
        sleep(1)
    if inp_num:
        if inp_num > len(results):
            print("out of index")
        for result in results:
            if result.index + 1 == inp_num:
                root = result
                break
    else:  # no specfic num was passed
        root = results[0]

    tree = []
    # because result.path isn't in in it's dirname, look down
    tree.append(root)
    for path in paths:
        if root.path in os.path.dirname(path.path):
            tree.append(path)
    return tree


def clean():
    '''Cleans the folder that contains the program from .docx and .pdf files'''
    count = 0
    for f in os.listdir(os.getcwd()):
        if '.pdf' in f or 'docx' in f:
            os.remove(f)
            count += 1
    print(f'Found {count} files.')
    sleep(2)


def create_pdf(results, inp, pdf_name='results',):
    '''Creates a pdf file from results that has the name passed to function '''
    def create_table(docx, results):
        '''creates a table in the docx file object and fills the info from
        results'''
        table = docx.add_table(rows=len(results)+1, cols=4)
        row = table.rows[0]
        row.cells[0].width = Inches(13)
        row.cells[0].text = 'Name'
        row.cells[1].width = Inches(0.1)
        row.cells[1].text = 'Size'
        row.cells[2].width = Inches(0.5)
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        row.cells[2].text = ''
        row.cells[3].text = 'Location'
        if type(results[-1]) == int:
            end = len(results) - 2
        else:
            end = len(results) - 1
        for index, result in enumerate(results[:end]):
            row = table.rows[index+1]
            row.cells[0].width = Inches(11)
            row.cells[0].text = result.name
            row.cells[1].width = Inches(0.1)
            size, unit = result.get_size()
            size = str(size)
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row.cells[1].text = f'{size:<3}'
            row.cells[2].width = Inches(0.5)
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            row.cells[2].text = f'{unit:<3}'
            row.cells[3].width = Inches(4)
            dir_tag = result.dir_tag.lower()
            if 'tv' in dir_tag.lower():
                dir_tag = dir_tag.replace('tv', 'TV')
            elif 'pc' in dir_tag.lower():
                dir_tag = dir_tag.replace('pc', 'PC')
            dir_tag = dir_tag.title()
            row.cells[3].text = dir_tag

    def create_docx(results, dest_path, docx_name='temp'):
        '''Takes in a list of results and the word file name(without .docx)
        and returns path of the created file'''
        print('Creating Docx...')
        docx_name += '.docx'
        docx = Document()
        for sec in docx.sections:
            sec.top_margin = Cm(0.5)
            sec.left_margin = Cm(0.01)
            sec.bottom_margin = Cm(0.5)
            sec.right_margin = Cm(0.5)
        print('Creating table.')
        create_table(docx, results)
        print('table created...')
        if dest_path:
            DOCX_PATH = f'{dest_path}{os.path.sep}{docx_name}'
        docx.save(DOCX_PATH)
        print('Docx file created succesfuly.')
        return DOCX_PATH

    match = re.search('export to .*', inp)
    if match:
        location = match.group().split(' to ')[1]
        if location == 'desk':
            dest_path = r'C:\Users\COMPU1\Desktop'
        else:
            dest_path = ''
    else:
        dest_path = os.getcwd()

    DOCX_PATH = create_docx(results, dest_path, pdf_name)
    print(f'Converting {DOCX_PATH} to PDF.')
    convert(DOCX_PATH)
    os.remove(DOCX_PATH)
    print(' .docx Converted succesfully.')
    sleep(2)


def is_rar(path):
    '''
    takes in a file path, returns true if it's a rar file or a r\d\d file'''
    match = re.search('(?:rar|r\d\d)$')
    if match:
        return True
    return False